"""Convert THESIS_REPORT.md into a submission-ready Word .docx that matches the
AIUB CSE thesis template (as exemplified by the department reference report).

Formatting: Times New Roman 12 / 1.5 spacing / justified body; centred title
page; chapter headings as two centred bold lines ("CHAPTER N" / TITLE) on a new
page; bold left section headings; centred italic figure/table captions; embedded
images; a running footer "CSC 4219 Project and Thesis [BSCS] .. Department of
Computer Science .. <page>"; lower-roman page numbers for the front matter and
decimal (restart at 1) for the body; and an auto-updating Table of Contents.

Run:  python -m src.md_to_docx
"""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parents[1]
MD = BASE / "THESIS_REPORT.md"
OUT = BASE / "THESIS_REPORT.docx"
MAX_W = 6.0
FOOTER_L = "CSC 4219 Project and Thesis [BSCS]"
FOOTER_C = "Department of Computer Science"

PAGEBREAK_PREFIXES = (
    "Declaration", "Approval", "Acknowledgement", "Author Contributions",
    "Project-Thesis Planning", "Table of Content", "List of Figures", "List of Tables",
    "List of Abbreviations", "Abstract", "References", "Appendix",
)


# ---------- low-level Word field / footer helpers -----------------------------
def _field(run, instr):
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instr
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, t, s, e):
        run._r.append(el)


def _set_pgnum(section, fmt, start=None):
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pg = OxmlElement("w:pgNumType"); pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def style_footer(section, fmt, start=None):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.text = ""
    p.paragraph_format.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(FOOTER_L + "\t" + FOOTER_C + "\t")
    _field(p.add_run(), "PAGE")
    for r in p.runs:
        r.font.size = Pt(9); r.font.name = "Times New Roman"
    _set_pgnum(section, fmt, start)


def add_toc_field(doc):
    p = doc.add_paragraph()
    _field_toc(p.add_run())


def _field_toc(run):
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve")
    t.text = 'TOC \\o "1-3" \\h \\z \\u'
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, t, s, hint, e):
        run._r.append(el)


# ---------- inline text -------------------------------------------------------
def add_runs(par, text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    for part in re.split(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        else:
            par.add_run(part)


def caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(10.5)


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5; pf.space_after = Pt(6); pf.space_before = Pt(0)

    sec0 = doc.sections[0]
    sec0.left_margin = sec0.right_margin = Inches(1.0)
    sec0.top_margin = sec0.bottom_margin = Inches(1.0)

    lines = MD.read_text(encoding="utf-8").splitlines()
    i = 0
    img_re = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")
    boldline_re = re.compile(r"^\*\*(.+)\*\*$")
    body_started = False

    while i < len(lines):
        line = lines[i].rstrip()

        # ---- title page block ----
        if line.strip() == "<<TITLEPAGE>>":
            i += 1
            first = True
            while i < len(lines) and lines[i].strip() != "<<END>>":
                raw = lines[i].rstrip()
                if raw:
                    tag, txt = (raw.split("|", 1) + ["", ""])[:2] if "|" in raw else ("INFO", raw)
                    txt = re.sub(r"`([^`]*)`", r"\1", txt)
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(txt)
                    if tag == "UNIV":
                        r.font.size = Pt(14); r.bold = False
                        p.paragraph_format.space_after = Pt(2)
                    elif tag == "LOGO":
                        r.font.size = Pt(11); r.italic = True; r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        p.paragraph_format.space_before = Pt(48); p.paragraph_format.space_after = Pt(48)
                    elif tag == "TITLE":
                        r.font.size = Pt(16); r.bold = True
                        p.paragraph_format.space_after = Pt(36)
                    elif tag == "NAME":
                        r.font.size = Pt(12)
                        p.paragraph_format.space_after = Pt(2)
                        if first:
                            p.paragraph_format.space_before = Pt(6); first = False
                    else:  # INFO
                        r.font.size = Pt(12)
                        p.paragraph_format.space_after = Pt(2)
                i += 1
            i += 1  # skip <<END>>
            continue

        if line.strip() in ("---", "") or line.startswith(">") or line.startswith("<<"):
            i += 1; continue

        if line.strip() == "{{TOC}}":
            add_toc_field(doc); i += 1; continue

        # ---- image + caption ----
        m = img_re.match(line)
        if m:
            cap_txt, path = m.group(1), m.group(2).strip()
            p = (BASE / path)
            if p.exists():
                pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.paragraph_format.space_before = Pt(6); pic.paragraph_format.space_after = Pt(2)
                try:
                    pic.add_run().add_picture(str(p), width=Inches(MAX_W))
                except Exception:
                    pic.add_run().add_picture(str(p))
                caption(doc, cap_txt)
            else:
                doc.add_paragraph(f"[MISSING FIGURE: {path}]")
            i += 1; continue

        # ---- table ----
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
            rows = [r for r in rows if not re.match(r"^[\s:\-]+$", "".join(r))]
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=ncol)
                t.style = "Light Grid Accent 1"; t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        cpar = t.rows[ri].cells[ci].paragraphs[0]
                        add_runs(cpar, row[ci] if ci < len(row) else "")
                        cpar.paragraph_format.space_after = Pt(2); cpar.paragraph_format.line_spacing = 1.0
                        for run in cpar.runs:
                            run.font.size = Pt(10)
                            if ri == 0:
                                run.bold = True
                sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6); sp.paragraph_format.line_spacing = 1.0
            continue

        # ---- headings ----
        hm = re.match(r"^(#{1,4})\s+(.*)$", line)
        if hm:
            level = len(hm.group(1)); txt = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", hm.group(2)).strip()
            if level == 1 and txt.startswith("CHAPTER"):
                if not body_started:
                    # front matter -> body: new section, restart decimal numbering
                    style_footer(sec0, "lowerRoman")
                    newsec = doc.add_section(WD_SECTION.NEW_PAGE)
                    newsec.left_margin = newsec.right_margin = Inches(1.0)
                    newsec.top_margin = newsec.bottom_margin = Inches(1.0)
                    style_footer(newsec, "decimal", start=1)
                    body_started = True
                else:
                    doc.add_page_break()
                parts = re.split(r"\s+—\s+|\s+-\s+", txt, maxsplit=1) if "—" in txt else [txt]
                for seg in parts:
                    hp = doc.add_paragraph(); hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    hp.paragraph_format.space_before = Pt(24 if seg is parts[0] else 0); hp.paragraph_format.space_after = Pt(6)
                    r = hp.add_run(seg.strip()); r.bold = True; r.font.size = Pt(18); r.font.name = "Times New Roman"
                i += 1; continue
            if level == 1 and any(txt.startswith(pre) for pre in PAGEBREAK_PREFIXES):
                if not body_started:
                    pass  # still front matter, just a page break
                doc.add_page_break()
            h = doc.add_heading("", level=min(level, 3))
            h.paragraph_format.space_before = Pt(10 if level <= 2 else 6); h.paragraph_format.space_after = Pt(4)
            szmap = {1: 15, 2: 13, 3: 12, 4: 12}
            r = h.add_run(txt); r.bold = True; r.font.size = Pt(szmap[level]); r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(0, 0, 0)
            i += 1; continue

        # ---- fully-bold standalone line (captions / mini-headings) ----
        bm = boldline_re.match(line.strip())
        if bm:
            inner = bm.group(1)
            if inner.startswith("Table ") or inner.startswith("Figure "):
                caption(doc, inner)
            else:
                p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
                p.add_run(inner).bold = True
            i += 1; continue

        # ---- lists ----
        if re.match(r"^\s*[-*]\s+", line):
            add_runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^\s*[-*]\s+", "", line)); i += 1; continue
        if re.match(r"^\s*\d+\.\s+", line):
            add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\s*\d+\.\s+", "", line)); i += 1; continue

        # ---- normal paragraph (justified) ----
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, line)
        i += 1

    if not body_started:
        style_footer(sec0, "lowerRoman")
    try:
        doc.save(OUT)
        print(f"saved -> {OUT}")
    except PermissionError:
        alt = OUT.with_name("THESIS_REPORT_FINAL.docx")
        doc.save(alt)
        print(f"'{OUT.name}' is open/locked; saved -> {alt}")


if __name__ == "__main__":
    main()
